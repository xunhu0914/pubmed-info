import pandas as pd
from docx import Document

# 读取包含文章信息的Excel文件
articles_file = '3D打印神经细胞相关文章.xlsx' #从PubMed2xl网址导入的文件
df_articles = pd.read_excel(articles_file, dtype={
    'Date of Publication': str,
    'Volume': str,
    'Issue': str,
    'Pages': str
})

# 读取包含期刊影响因子的Excel文件
impact_factors_file = 'JCR2023-UTF8.xlsx' #2023年JCR分数
df_impact_factors = pd.read_excel(impact_factors_file)

# 将期刊名称转换为小写以便不区分大小写匹配
df_articles['Journal Title'] = df_articles['Journal Title'].str.lower()
df_impact_factors['Journal'] = df_impact_factors['Journal'].str.lower()
#待解决：有的期刊名称中含有（）导致无法匹配

# 合并两个DataFrame，使用左连接
merged_df = pd.merge(df_articles, df_impact_factors, left_on='Journal Title', right_on='Journal', how='left')

# 删除重复的期刊名称列
merged_df.drop(columns=['Journal'], inplace=True)

# 导出合并后的Excel文件
merged_file = 'Merged_Articles_with_IF.xlsx'
merged_df.to_excel(merged_file, index=False)

# 添加一个断点，以便我能手动补全缺失的影响因子
import pdb; pdb.set_trace()

# 重新读取手动修改后的Excel文件
merged_df = pd.read_excel(merged_file)

# 创建一个新的Word文档
doc = Document()

# 遍历合并后的DataFrame中的每一行
for index, row in merged_df.iterrows():
    # 提取标题、作者、摘要和影响因子
    title = row['Title']
    authors = row['Author(s) Full Name']
    abstract = row['Abstract']
    impact_factor = row['IF']
    date_of_publication = row['Date of Publication'][:4] if pd.notna(row['Date of Publication']) else ''
    volume = '' if pd.isna(row['Volume']) else f"Vol. {row['Volume']}"
    issue = '' if pd.isna(row['Issue']) else f"Issue. {row['Issue']}"
    pages = '' if pd.isna(row['Pages']) else f"pp. {row['Pages']}"

    # 构建影响因子字符串
    impact_factor_str = f' (IF: {impact_factor})' if pd.notna(impact_factor) else ''
    
    #构建出版信息字符串
    publication_info_parts = f"{date_of_publication} {volume} {issue} {pages}".strip()
    
    # 将信息添加到Word文档
    doc.add_paragraph(f'{index + 1}. {title} {impact_factor_str} {authors}  {publication_info_parts} ')
    doc.add_paragraph(abstract)
    doc.add_page_break()  # 每个摘要后添加分页符

# 保存Word文档
word_file = 'Extracted_Articles_with_IF.docx'
doc.save(word_file)



