import pandas as pd
from docx import Document
import json
import tencentcloud.common as common
from tencentcloud.common.exception.tencent_cloud_sdk_exception import TencentCloudSDKException
from tencentcloud.tmt.v20180321 import tmt_client, models

# 腾讯云API密钥
SecretId = "your_secret_id"  # 替换为你的SecretId
SecretKey = "your_secret_key"  # 替换为你的SecretKey

# 初始化腾讯云翻译客户端
cred = common.Credential(SecretId, SecretKey)
client = tmt_client.TmtClient(cred, "ap-guangzhou")  # 根据你的区域选择正确的ID

def translate_text(text, source_lang='zh', target_lang='en'):
    try:
        # 创建请求对象
        req = models.TextTranslateRequest()
        params = {
            "SourceText": text,
            "Source": source_lang,  # 语言源代码，zh为中文
            "Target": target_lang,  # 目标语言，en为英文
            "ProjectId": 0  # 使用默认的项目ID
        }
        req.from_json_string(json.dumps(params))
        
        # 发起请求
        response = client.TextTranslate(req)
        
        # 解析响应
        result = response.to_json_string()
        response_json = json.loads(result)
        translated_text = response_json['TargetText']
        return translated_text
    
    except TencentCloudSDKException as err:
        print(f"API请求失败: {err}")
        return None

# 读取Excel文件
excel_file = '3D打印神经细胞相关文章.xlsx'
df = pd.read_excel(excel_file)

# 创建一个新的Word文档
doc = Document()

# 遍历DataFrame中的每一行
for index, row in df.iterrows():
    # 提取标题、作者和摘要
    title = row['Title']
    authors = row['Author(s) Full Name']
    abstract = row['Abstract']
    
    # 翻译摘要
    translated_abstract = translate_text(abstract, source_lang='zh', target_lang='en')  # 根据需要调整语言
    
    # 将信息添加到Word文档
    doc.add_paragraph(f'{index + 1}. {title} {authors}')
    doc.add_paragraph(f'原文摘要: {abstract}')
    doc.add_paragraph(f'翻译摘要: {translated_abstract if translated_abstract else "翻译失败"}')
    doc.add_page_break()  # 每个摘要后添加分页符

# 保存Word文档
word_file = 'Extracted_Articles_with_Translation.docx'
doc.save(word_file)
