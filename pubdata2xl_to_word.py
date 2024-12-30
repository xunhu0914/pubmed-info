#确保pandas,python-docx,openpyxl在环境中正确的安装
#确保本程序与excel处于同一路径下
import pandas as pd
from docx import Document

# 读取Excel文件
excel_file = '3D打印神经细胞相关文章.xlsx'
df = pd.read_excel(excel_file)

# 创建一个新的Word文档
doc = Document()

# 遍历DataFrame中的每一行
for index, row in df.iterrows():
    # 提取标题、作者和摘要
    title = row['Title']
    authors = row['Author(s) Full Name']
    abstract = row['Abstract']
    
    # 将信息添加到Word文档
    doc.add_paragraph(f'{index + 1}. {title} {authors}')
    doc.add_paragraph(abstract)
    doc.add_page_break()  # 每个摘要后添加分页符

# 保存Word文档
word_file = 'Extracted_Articles.docx'
doc.save(word_file)